import streamlit as st
import numpy as np
import matplotlib as mpl
import tempfile
import os

from helpers import *

from docx import Document
from io import BytesIO
from datetime import datetime

from openai import OpenAI
from google import genai
from google.genai import types

def concatenate_input():
    full_listing = ""

    for i in range(1, 6):
        subheading = st.session_state.get(f"subheading{i}", "")
        bullet = st.session_state.get(f"bullet{i}", "")
        
        complete_bullet = subheading + ":" + bullet + "\n"
        full_listing += complete_bullet + "\n"
        
    return full_listing

def extract_paragraphs_by_phrase(text: str, phrase: str):
    """
    Return a list of all paragraphs/lines that contain `phrase`.
    A "paragraph" is defined as the text from a line start (or previous newline)
    up to the next newline. Matching is case-insensitive.
    """
    if not phrase:
        return []

    # escape phrase (so user input like "vintage (design)" won't break the regex)
    esc = re.escape(phrase)
    # pattern: start at line boundary (^ or \n), capture any chars except newline that include the phrase
    pattern = rf'(?im)(?:^|\n)([^\n]*\b{esc}\b[^\n]*)'
    matches = re.findall(pattern, text)
    # strip extra whitespace and return
    return [m.strip() for m in matches]


def extract_first_paragraph_by_phrase(text: str, phrase: str):
    """
    Return the first paragraph/line that contains `phrase` or None if not found.
    """
    res = extract_paragraphs_by_phrase(text, phrase)
    return res[0] if res else None

from PIL import Image, UnidentifiedImageError, ImageOps

def colorize_df(df: pd.DataFrame, base_cmap: str = "Greens", min_frac: float = 0.0, max_frac: float = 0.8):
    # cmap_orig = mpl.cm.get_cmap(base_cmap)
    cmap_orig = mpl.colormaps[base_cmap]
    
    cmap_vals = cmap_orig(np.linspace(min_frac, max_frac, 256))
    cmap_vals = cmap_vals[:, :-1]
    soft_cmap = mpl.colors.ListedColormap(cmap_vals)

    return (
        df.style
        .background_gradient(cmap=soft_cmap, vmin=0, vmax=100, axis=None)
        .format("{:.0f}")
    )

def get_cluster_dict(picture_cluster_df_labeled):
    """
    Create a dictionary mapping each cluster label to a list of unique keywords sorted alphabetically.
    """
    cluster_dict = picture_cluster_df_labeled.groupby("label")["keywords"].apply(list).to_dict()
    
    # Flatten all keyword lists per label, remove duplicates, and sort
    cluster_dict = {
        k: sorted(set([item for sublist in v for item in sublist]))
        for k, v in cluster_dict.items()
    }
    
    return cluster_dict

def total_sales_for_phrase(df, phrase: str, col='parsed_headings', sales_col='monthly_sales'):
    """
    Sum monthly sales for all rows where at least one element in df[col]
    (a List[str]) contains the given phrase (case-insensitive).
    """
    phrase = phrase.lower()

    mask = df[col].apply(
        lambda lst: any(
            phrase in str(item).lower() 
            for item in lst if pd.notna(item)
        )
    )

    return df.loc[mask, sales_col].sum()

def total_sales_for_phrases(df, phrases: list, col='parsed_headings', sales_col='monthly_sales'):
    """
    For each phrase in phrases, compute total sales using the modular
    total_sales_for_phrase() function. Return a ranked DataFrame.
    """
    results = []

    for phrase in phrases:
        total = total_sales_for_phrase(df, phrase, col=col, sales_col=sales_col)
        results.append({"phrase": phrase, "total_sales": total})

    # Convert to DataFrame and sort
    result_df = pd.DataFrame(results).sort_values(
        by="total_sales", ascending=False
    ).reset_index(drop=True)

    result_df['total_sales'] = result_df['total_sales'].astype(int)

    return result_df

def filter_bullets_by_phrase(df, phrase: str, col='bullet_points'):
    """
    Return rows where df[col] (a string) contains the phrase
    as a substring (case-insensitive).
    """
    phrase = phrase.lower()

    mask = df[col].fillna("").str.lower().str.contains(phrase)

    return df[mask]

# --- Streamlit Setup ---
st.set_page_config(page_title="Image Keyword Filter", layout="wide")
st.title("Amazon Listing Dashboard")

if "label_dict" not in st.session_state:
    st.session_state['bullet_clusters'] = pd.read_parquet('antique_bullet_clusters.parquet')
    st.session_state['label_dict'] = get_cluster_dict(st.session_state['bullet_clusters'])

label_dict = st.session_state['label_dict']

if "bullet_diagram" not in st.session_state:
    st.session_state['bullet_diagram'] = pd.read_parquet('bullet_diagram.parquet')
    st.session_state['labels'] = st.session_state['bullet_diagram'].index.tolist()

labels = st.session_state['labels']
bullet_diagram = st.session_state['bullet_diagram']
bullet_diagram = colorize_df(bullet_diagram)


if "bullet_labels" not in st.session_state:
    st.session_state['bullet_labels'] = pd.read_parquet('bullet_labels.parquet')

bullet_labels = st.session_state['bullet_labels']

if "title_keywords" not in st.session_state:
    st.session_state['title_keywords'] = pd.read_parquet('antique_picture_frame_keywords_final.parquet')

title_keywords = st.session_state['title_keywords']
 
if "best_phrases" not in st.session_state:
    st.session_state['best_phrases'] = {k: total_sales_for_phrases(bullet_labels, label_dict[k]) for k in labels}

best_phrases = st.session_state['best_phrases']


left_col, _, right_col = st.columns([7.8, .8, 8])



with left_col:
    st.write("\n")
    st.dataframe(bullet_diagram, use_container_width=True)

with right_col:
    selected_label = st.selectbox("Select a category:", labels)

    st.markdown(f"#### {selected_label}")
    st.dataframe(best_phrases[selected_label], use_container_width=True, height=248)


def update_filter():
    label = st.session_state["label_input"]
    st.session_state["filtered"] = filter_bullets_by_phrase(
        bullet_labels, label
    )

# Text input with on-change trigger
with st.expander("Search for examples"):

    st.text_input(
        "Enter Label",
        key="label_input",
        on_change=update_filter,
        width = 250
    )

    # # Display filtered results (if any)
    if "filtered" in st.session_state:
        filtered_df = st.session_state["filtered"]
        # st.write(filtered_df.head())

    DISPLAY_NUMBER = 50

    if st.button("🎨 Load Images"):
        # If more than 24 images, randomly select 24 rows
        if len(filtered_df) > DISPLAY_NUMBER:
            trimmed_sample = filtered_df.head(DISPLAY_NUMBER)
        else:
            trimmed_sample = filtered_df.copy()

        # Display images in 3-column grid
        
        def to_str(val):
            """Convert list/ndarray to readable comma-separated string."""
            if isinstance(val, (list, set, tuple)):
                return ", ".join(map(str, val))
            if isinstance(val, np.ndarray):
                return ", ".join(map(str, val.tolist()))
            return str(val)

        # FIXED_SIZE = (1200, 1200)  # width, height for all images
        grid_cols = st.columns(3)


        for idx, (_, row) in enumerate(trimmed_sample.iterrows()):
            with grid_cols[idx % 3]:
                img_path = row["image_path"]
                try:
                    # Display image
                    st.image(img_path)
                    
                    # url
                    url = row['url']
                    st.caption(url)

                    # Extract bullet
                    full_listing = to_str(row.get("bullet_points", ""))
                    specific_bullet = extract_first_paragraph_by_phrase(full_listing, st.session_state["label_input"])


                    st.markdown(
                        f"""
                        Monthly Sales: {to_str(row.get("monthly_sales", []))}  
                        """
                    )

                    with st.expander('Show full listing'):
                        # st.text(row['bullet_points'])
                        st.markdown(row['bullet_points'].replace("\n", "\n\n"))
                    # Force a fixed height container for text
                    st.markdown(
                        f"<div style='height:240px; overflow:hidden'>{specific_bullet}</div>",
                        unsafe_allow_html=True
                    )

                

                except (FileNotFoundError, UnidentifiedImageError, OSError):
                    st.warning(f"⚠️ Could not load image: {img_path}")


# NOTES: Instead of total sales, maybe also do number of listings who use its

# ---------------------------------------
#           OPENAI PHOTO FEATURE
# ---------------------------------------

@st.cache_resource
def get_openai_client(api_key):
    if "openai_client" not in st.session_state:
        st.session_state.openai_client = OpenAI(api_key=api_key)
    return st.session_state.openai_client

client = get_openai_client(st.secrets["OPENAI_KEY"])

@st.cache_resource
def get_gemini_client(api_key):
    if "gemini_client" not in st.session_state:
        st.session_state.gemini_client = genai.Client(api_key=api_key)
    return st.session_state.gemini_client

gemini_client = get_gemini_client(st.secrets["GEMINI_KEY"])

st.session_state.setdefault("photo_result", None)
st.session_state.setdefault("title_result", None)
st.session_state.setdefault("suggestion_result", None)
st.session_state.setdefault("synonym_result", None)
st.session_state.setdefault("product_components_result", None)
st.session_state.setdefault("rewriting_result", None)
st.session_state.setdefault("previous_user_input", None)
st.session_state.setdefault("previous_user_phrase", None)
st.session_state.setdefault("previous_user_input_synonym", None)
st.session_state.setdefault("ai_expander", False)
st.session_state.setdefault("default_tab", "Keywords & Title")

st.session_state.setdefault("pre_optimized_listing", False)

st.session_state.setdefault("finished_product_title", "")


# Listing optimizer
st.session_state.setdefault("input_listing", None)
st.session_state.setdefault("optimized_listing", None)

image_description_col, _, ai_tools_col = st.columns([5, 1, 8])

# =====================================================
# Image Description Generator
# =====================================================
with image_description_col:
    st.markdown("#### Amazon Writing Assistant")

    @st.cache_data(show_spinner=False)
    def load_and_resize_image(file, max_size=1024):
        image = Image.open(file)
        image.thumbnail((max_size, max_size))
        
        # reset description
        st.session_state["photo_result"] = ""
        return image

    uploaded_image = st.file_uploader(
        "Upload a product photo",
        type=["jpg", "jpeg", "png"],
        key="uploaded_image"
    )

    def generate_description_callback():
        if uploaded_image is not None:
            image = load_and_resize_image(uploaded_image)
            with st.spinner("Generating description..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                    if image.mode in ("RGBA", "P"):
                        image = image.convert("RGB")

                    image.save(tmp.name, format="JPEG")
                    temp_image_path = tmp.name

                try:
                    st.session_state["photo_result"] = generate_product_description_from_image(
                        client=client,
                        image_file=temp_image_path,
                        system_instructions=image_describe_system_instructions,
                        processing_instructions=image_describe_processing_instructions,
                    )
                finally:
                    os.remove(temp_image_path)

    if uploaded_image is not None:
        image = load_and_resize_image(uploaded_image)
        st.image(image, caption="Uploaded Image", use_container_width=True)

        # Button using on_click
        if not st.session_state["photo_result"]:
            st.button(
                "Generate description",
                on_click=generate_description_callback
            )

        # ---- Persistent render ----
        if st.session_state["photo_result"] is not None:
            st.markdown(st.session_state["photo_result"])

    st.write("")
    with st.expander("Product details"):
        facts_col, _, keyword_col = st.columns([8, 0.5, 6])
        with facts_col:
            st.text_area("Specs", height=200, key='item_specifications')
        with keyword_col:
            st.text_area("Keywords", height=200, key='listing_bullet_keywords')
            keywords_list = re.split(r"[,\s]+", st.session_state['listing_bullet_keywords'].strip())

        if st.button("Generate listing"):
            # Step 1: Listing 1st draft 
            listing_prompt =(
                listing_writer_instructions + st.session_state['item_specifications'] + "\n\n"
                + 'Keywords: ' + ", ".join(st.session_state['listing_bullet_keywords'].split()) + "\n\n" + 
                listing_writer_example + "\n\n" + "Your output:"

            )

            with st.spinner("Generating listing..."):
                st.session_state["pre_optimized_listing"] = gemini_client.models.generate_content(
                        model="gemini-3-pro-preview",
                        contents=listing_prompt,
                        config=types.GenerateContentConfig(
                            thinking_config=types.ThinkingConfig(thinking_level="low")
                        ),
                    ).text
                
            # Step 2: Optimize listing with keywords
            optimize_listing_prompt = (
                keyword_enrich_header_instructions + 
                st.session_state['listing_bullet_keywords'] + "\n\n"
                "Listing Input:\n\n" + 
                st.session_state['pre_optimized_listing'] + 
                "Output:"
            )

            with st.spinner("Optimizing listing..."):
                st.session_state["optimized_listing"] = gemini_client.models.generate_content(
                            model="gemini-3-pro-preview",
                            contents=optimize_listing_prompt,
                            config=types.GenerateContentConfig(
                                thinking_config=types.ThinkingConfig(thinking_level="low")
                            ),
                        ).text

    # Display results!
    def keyword_markdown(title, keyword_set):
        if not keyword_set:
            st.markdown(f"**{title}:** None ✅")
            return

        kw_string = ", ".join(sorted(keyword_set))
        st.markdown(f"**{title}:** {kw_string}")

    def display_listing_interface(listing, keyword_count_df, all_keywords, total_keywords_unique, new_words_added = None):
            st.write("")
            st.write(listing) 
            st.divider()

            stats_col1, _, stats_col2 = st.columns([7, 3, 5])
            with stats_col1:
                
                st.write(f"**Keywords provided**: {all_keywords}") 
                st.write(f"**Unique keywords in listing**: {total_keywords_unique}")
                if new_words_added:
                    keyword_markdown("New keywords added", new_words_added)
            
            with stats_col2:
                st.dataframe(keyword_count_df, hide_index=True, height = 176, width = 200)

    if st.session_state["pre_optimized_listing"] and st.session_state["optimized_listing"]:
        # Aggregate info
        listing_before = st.session_state["pre_optimized_listing"]
        listing_after = st.session_state["optimized_listing"]
        
        all_keywords = len(keywords_list)
        total_keywords_unique1, df1 = summarize_listing_keyword_stats(keywords_list, listing_before)
        total_keywords_unique2, df2 = summarize_listing_keyword_stats(keywords_list, listing_after)

        list1 = set(df1.keyword.unique().tolist())
        list2 = set(df2.keyword.unique().tolist())
        new_keywords_added = list2 - list1
        not_included = set(keywords_list)-list2

        # Display Info
        st.toggle("SEO-optimized", key = 'optimized')
        if not st.session_state['optimized']:
            display_listing_interface(listing_before, df1, all_keywords, total_keywords_unique1)
        else:
            display_listing_interface(listing_after, df2, all_keywords, total_keywords_unique2, new_keywords_added)
       
# =====================================================
# Keywords / Writing Assistant
# =====================================================
with ai_tools_col:
    st.markdown("##")
    with st.expander("AI writing tools", expanded = st.session_state['ai_expander']):
        tab1, tab2, tab3 = st.tabs(["Keywords & Title", "Proofreader", "Writing Assistant"], default=st.session_state["default_tab"])

        with tab1:
            col1, _, col2, _ = st.columns([4.2, 0.36, 5, 0.1])

            with col1:
                st.markdown("##### Title generator")
                st.text_area(
                    "Keywords",
                    placeholder="e.g. picture, frame, gold, ornate",
                    key="input_keywords",
                    height=280
                )

                def generate_product_title():
                    user_input = st.session_state.get("input_keywords", "").strip()

                    if not user_input:
                        st.session_state["title_result"] = ""
                        return

                    title_prompt = generate_title_prompt_prefix + user_input + "Output:"
                    st.session_state["title_result"] = gemini_client.models.generate_content(
                        model="gemini-3-flash-preview",
                        contents=title_prompt,
                        config=types.GenerateContentConfig(
                            thinking_config=types.ThinkingConfig(thinking_level="low")
                        ),
                    ).text

                    # st.session_state["title_result"] = gemini_client.models.generate_content(
                    #     model="gemini-2.5-flash-lite",
                    #     contents=title_prompt
                    # ).text

                    # model gemini-3-flash-preview
                
                st.write("")
             

                if st.button("Generate title"):
                    with st.spinner("Generating title..."):
                        generate_product_title()

                if st.session_state["title_result"]:
                    st.write("")
                    st.markdown(f"**{st.session_state["title_result"]}**")
            # ---- Persistent render ----
            
            with col2:
                st.markdown("##### Keyword list")
                st.dataframe(title_keywords, height = 352, hide_index=True)
                st.text_area("Final title", key = 'finished_product_title')

            with tab2:
                col1, _, col2, _ = st.columns([4.2, 0.2, 5, 0.1])

                def save_user_phrase():
                    st.session_state['previous_user_phrase'] = st.session_state['input_to_rephrase'].strip()  
                    st.session_state['ai_expander'] = True
                    st.session_state["default_tab"] = "Proofreader"

                with col1:
                    st.text_area(
                        "Enter any phrase or sentence:",
                        placeholder="e.g. Cleaning the glass surface soft cloth provided to remove finger mark if any",
                        key="input_to_rephrase",   # 🔑 persistent key
                        value=st.session_state['previous_user_phrase'],
                        on_change=save_user_phrase,
                        max_chars=200,
                        width=400,
                        height=200
                    )

                    if st.button("Rewrite phrase"):
                        user_prompt = (
                            amazonify_prompt_prefix
                            + st.session_state["input_to_rephrase"].strip()
                            + "Output:"
                        )

                        with st.spinner("Rewriting phrase..."):
                            st.session_state["rewriting_result"] = complete_phrase(
                                client,
                                user_prompt,
                                model='gpt-5.1'
                            )

                # ---- Persistent render ----
                with col2:
                    if st.session_state["rewriting_result"]:
                            st.write("")
                            st.markdown("**Alternative phrasing**")
                            st.write(st.session_state["rewriting_result"])
                    
                    # if st.session_state["product_components_result"]:
                    #     st.write("")
                    #     st.markdown("**Product components**")
                    #     st.write(st.session_state["product_components_result"])
            
            with tab3:
                col1, _, col2, _ = st.columns([4.2, 0.2, 5, 0.1])

                if "mode_rewrite" not in st.session_state:
                    st.session_state["mode_rewrite"] = False

                def toggle_mode():
                    st.session_state["mode_rewrite"] = not st.session_state["mode_rewrite"]

                def save_user_input():
                    st.session_state['previous_user_input'] = st.session_state['user_phrase'].strip()  
                    st.session_state['ai_expander'] = True
                    st.session_state["default_tab"] = "Writing Assistant"

                def save_user_input_2():
                    st.session_state['previous_user_input_synonym'] = st.session_state['user_phrase_for_synonym'].strip()  
                    st.session_state['ai_expander'] = True
                    st.session_state["default_tab"] = "Writing Assistant"

                with col1:
                    st.toggle(
                        "Rewrite mode", 
                        value=st.session_state["mode_rewrite"],
                        on_change=toggle_mode,
                        help="Switch between generating suggestions and rewriting a phrase"
                    )

                    st.write("")
                    # ===============================
                    # MODE 1 — GENERATE SUGGESTIONS
                    # ===============================
                    if not st.session_state["mode_rewrite"]:
                        st.text_input(
                            "Enter any phrase:",
                            placeholder="e.g. sculpted floral accents",
                            key="user_phrase",  
                            value = st.session_state['previous_user_input'],
                            on_change=save_user_input,
                            width=400
                        )

                        if st.button("Generate suggestions"):
                            user_prompt = (
                                generate_suggestion_prompt_prefix
                                + st.session_state["user_phrase"].strip()
                                + generate_suggestions_prompt_suffix
                            )

                            with st.spinner("Generating suggestions..."):
                                st.session_state["suggestion_result"] = complete_phrase(
                                    client,
                                    user_prompt
                                )

                    # ===============================
                    # MODE 2 — REWRITE PHRASE
                    # ===============================
                    else:
                        st.text_input(
                            "Enter any phrase:",
                            placeholder="e.g. makes changing photos easy",
                            key="user_phrase_for_synonym",  
                            value = st.session_state['previous_user_input_synonym'],
                            on_change=save_user_input_2,
                            max_chars=60,
                            width=400
                        )

                        if st.button("Generate synonyms"):
                            user_prompt = (
                                generate_synonyms_prompt_prefix
                                + st.session_state["user_phrase_for_synonym"].strip()
                                + "Output:"
                            )

                            with st.spinner("Generating synoynms..."):
                                st.session_state["synonym_result"] = complete_phrase(
                                    client,
                                    user_prompt
                                )

                # ---- OUTPUT COLUMN ----
                with col2:
                    if not st.session_state["mode_rewrite"]:
                        if st.session_state["suggestion_result"]:
                            st.markdown("**Generated Suggestions**")
                            st.write(st.session_state["suggestion_result"])
                    else:
                        if st.session_state["synonym_result"]:
                            st.markdown("**Synonyms**")
                            st.write(st.session_state["synonym_result"])
                            

    st.write("")
    st.write("")
    about_col, _, title_col = st.columns([3.6, 0.2, 10])
    with title_col:
        st.markdown(f"##### {st.session_state["finished_product_title"]}")
        st.write("")
    
    st.write("")
    for i in range(1, 6):
        subheading_col, _, bullet_col = st.columns([3.6, 0.2, 10])

        with subheading_col:
            st.text_input(
                "Subheading",
                key=f"subheading{i}",
                width=190
            )

        with bullet_col:
            st.text_area(
                "Content",
                key=f"bullet{i}",
                height=128,
                max_chars=300

            )

    # Download Button!
    def build_listing_docx(listing_text: str, product_title: str) -> BytesIO:
        doc = Document()

        # ---- PRODUCT TITLE ----
        if product_title and len(product_title) > 10:
            doc.add_heading(product_title, level = 1)
        else:
            doc.add_heading("TITLE MISSING", level = 1)
        
        doc.add_paragraph("")  # spacing

        # preserve line breaks
        for line in listing_text.split("\n"):
            doc.add_paragraph(line)

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        # Save to memory
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    st.session_state['input_listing'] = concatenate_input()    
    listing_text = st.session_state.get("input_listing", "")
    product_title = st.session_state.get("finished_product_title", "")
    st.write("")

    # if not listing_text or len(listing_text) < 100:
    #     st.warning("Please fill in Amazon listing")
    # elif not product_title or len(product_title) < 20:
    #     st.warning("Please write a product title")
    if listing_text and product_title:
        docx_file = build_listing_docx(listing_text, product_title)
        st.download_button(
            label="Download listing",
            data=docx_file,
            file_name="amazon_listing.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )